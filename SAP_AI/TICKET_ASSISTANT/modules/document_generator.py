"""Generate the ticket evidence Word document."""
from datetime import datetime
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


class DocumentGenerator:
    def __init__(self, organization_name="Your Organization"):
        self.organization_name = organization_name

    def generate(self, ticket, steps, audit_events, output_path):
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc = Document()
        section = doc.sections[0]
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)

        title = doc.add_heading(f"{ticket['ticket_id']} Analysis Report", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Organization: {self.organization_name}")
        doc.add_paragraph(f"Ticket Type: {ticket['issue_type']}")
        doc.add_paragraph(f"Status: {ticket['status']}")
        doc.add_paragraph(f"Created: {ticket['created_at']}")
        doc.add_paragraph(f"Generated: {datetime.now().astimezone().isoformat(timespec='seconds')}")
        doc.add_paragraph("Privacy notice: This report is intended to contain sanitized technical evidence only. It must not contain production business records or personal data.")

        doc.add_heading("Resolution Evidence", level=1)
        for step in steps:
            doc.add_heading(f"Step {step['step_number']}: {step['step_name']}", level=2)
            doc.add_paragraph(f"Status: {step['status']}")
            doc.add_paragraph(f"Completed: {step['completed_at'] or 'Not completed'}")
            if step.get('ideas'):
                doc.add_paragraph("Checks performed:")
                for idea in step['ideas']:
                    doc.add_paragraph(idea, style='List Bullet')
            image_path = step.get('screenshot_path')
            if image_path and Path(image_path).exists():
                doc.add_picture(image_path, width=Inches(6.1))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                doc.add_paragraph(f"Evidence file: {Path(image_path).name}")
            else:
                doc.add_paragraph("Evidence image unavailable.")

        doc.add_heading("Audit Trail", level=1)
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Timestamp", "Event", "Details"
        for event in audit_events:
            cells = table.add_row().cells
            cells[0].text = event['occurred_at']
            cells[1].text = event['event_type']
            cells[2].text = event['details']
        doc.save(output_path)
        return output_path
